# Aula 6: Criação e edição de documentos Word com python-docx
# Duração: ~1h30

from docx import Document
from docx.shared import Inches

# Parte 1: Criar documento Word
doc = Document()
doc.add_heading("Relatório de Livros", 0)

livros=[]
for livro in livros:
    doc.add_heading(livro["Título"], level=1)
    doc.add_paragraph(f"Autor: {livro['Autor']}")
    doc.add_paragraph(f"Ano: {livro['Ano']}")

# Parte 2: Inserir uma tabela
tabela = doc.add_table(rows=1, cols=3)
header_cells = tabela.rows[0].cells
header_cells[0].text = 'Título'
header_cells[1].text = 'Autor'
header_cells[2].text = 'Ano'

for livro in livros:
    row_cells = tabela.add_row().cells
    row_cells[0].text = livro['Título']
    row_cells[1].text = livro['Autor']
    row_cells[2].text = str(livro['Ano'])

# Salvar documento
doc.save("relatorio_livros.docx")
print("Documento Word criado com sucesso.")