# pip install python-docx ->  Biblioteca para trabalhar com arquivos Word

from docx import Document

#1 - Criar  um novo documento
documento = Document()

#2 - Adicionar um Título principal
documento.add_heading('Relatório de Atividades', level=1)

#3 - Adicionar parágrafo com texto
documento.add_paragraph('Este documento foi gerado durante a aula de RPA')
documento.add_paragraph('Aqui colocaria um novo texto...')

#4- Adicionar um subtítulo
documento.add_heading('Atividades', level=2)

#5- Adicionar a lista de Atividades
atividades = [
    'Reunião com a equipe de projetos',
    'Desenvolvimento de novos módulos',
    'Testes de funcionalidade',
    'Treinamento para novos colaboradores'
]
for atividade in atividades:
    documento.add_paragraph(atividade,style='List Bullet')

#6 Novo Subtítulo de considerações finais
documento.add_heading('Considerações Finais', level=2)

#7 Parágrafo final
documento.add_paragraph('Todas as metas estabelecidas foram atingidas.')

#8 Salvar o arquivo
documento.save('Relatorio_atividades.docx')
print("Arquivo criado com sucesso!")

