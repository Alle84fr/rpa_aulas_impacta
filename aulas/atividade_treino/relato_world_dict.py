# Atividade prática : Coleta e relatório Automatizado de dados Geográficos com Python

# Entrega esperada
# - Código-fonte em Python (.py);
# - Banco de dados .db contendo os registros;
# - Arquivo de relatório final (.xlsx ou .docx).

# ALessandra Furlanetto Rigonatti 2401151 - ADS - Manhã - A
# Queria deixar o relatório mais bonito, tive um pouco de dificuldade, preciso estudar mais relatório, é a minha pedra

import requests
import sqlite3
import sys
import io
from docx import Document
from docx.shared import Pt

# ___________________ CARACTERES ESPECIAIS ____________________________________

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# __________________________ DEF RELATÓRIO ____________________________________

def relatorio (paises):
    ''' Cria relatório, no formato World (.docx)'''

    doc = Document()
    
    style = doc.styles["Normal"] #estilo básico, padrão do dock
    font_config = style.font 
    font_config.name = "Arial"
    font_config.size = Pt(12)
    
    titulo = doc.add_paragraph()
    titulo.alignment = 1    #centraliza o título
    titulo.add_run("🌍 Relatório 3 Países 🌏 ").bold = True
    

    # ------------------- Listando dados de cada país -----------------------------
    
    for pais in paises:
                
        doc.add_paragraph(f"\n- {pais['Nome']}", style='List Bullet') # list Bullet = forma uma lista com bolinhas, que irá separar a lista dos países
        
        doc.add_paragraph(f"Nome Oficial: {pais['Nome_Oficial']}")
        doc.add_paragraph(f"Capital: {pais['Capital']}")
        doc.add_paragraph(f"Continente: {pais['Continente']}")
        doc.add_paragraph(f"Região: {pais['Regiao']}")
        doc.add_paragraph(f"Sub-região: {pais['Sub_Regiao']}")
        doc.add_paragraph(f"População: {pais['Populacao']}")
        doc.add_paragraph(f"Área: {pais['Area']}")
        doc.add_paragraph(f"Fuso Horário: {pais['Fuso_Horario']}")
        doc.add_paragraph(f"Moeda: {pais['Simbolo']} {pais['Nome_Moeda']}")
        doc.add_paragraph(f"Idioma: {pais['Idioma']}")
        
        doc.save("paises.docx")

# ____________________________ DADOS DO URL ____________________________________

n = 3
paises = []

while n > 0:

    pais = input("🌎 Nome de um país, em inglês 👈: ")

    url = f"https://restcountries.com/v3.1/name/{pais}"

    resposta = requests.get(url)

    if resposta.status_code == 200:
        dados = resposta.json()
        # print(resposta.status_code)
    else:
        print("Error:", f"Erro na requisição: {resposta.status_code}")
        
   # ------------------- Selecionando dados de cada país -----------------------------
   
    dictotal = dados[0]
    nome = dictotal["name"]["common"]
    nomeOffi = dictotal["name"]["official"] 
    capital = dictotal["capital"][0] if "capital" in dictotal else "N/A"
    # se capital no dict mostre, se não capital marque como Not Available - Não disponível, não applicável
    continente = dictotal["continents"][0] if "continents" in dictotal else "N/A"
    regiao = dictotal["region"] if "region" in dictotal else "N/A"
    subRegiao = dictotal["subregion"]if "subregion" in dictotal else "N/A"
    populacao = dictotal["population"] if "population" in dictotal else "N/A"
    area = dictotal["area"] if "area" in dictotal else "N/A"
    fusoH = ", ". join(dictotal["timezones"]) if "timezones" in dictotal else "N/A" #aqui retorna lista, tive que converter para string
    bandeira = dictotal["flags"]["png"] if "flags" in dictotal and "png" in dictotal["flags"] else "N/A"

    # modeda é lista não dicionário então
    moeda = list(dictotal["currencies"].values())[0] if "currencies" in dictotal else {"symbol": "N/A", "name": "N/A"}
    simbo = moeda["symbol"]
    nomeMoed = moeda["name"]

    idioma = list(dictotal["languages"].values())[0] if "languages" in dictotal else "N/A"
    

    # ------------------- Criando arquivo -----------------------------
    
    conexao = sqlite3.connect("paises.bd")
    cursor = conexao.cursor() 
    cursor.execute (""" CREATE TABLE IF NOT EXISTS paises(
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        Nome TEXT,
                        Nome_Oficial TEXT,
                        Capital TEXT,
                        Continente TEXT,
                        Regiao TEXT,
                        Sub_Regiao TEXT,
                        Populacao INTEGER,
                        Area REAL,
                        Fuso_Horario TEXT,
                        Bandeira TEXT,
                        Nome_Moeda TEXT,
                        Simbolo TEXT,
                        Idioma TEXT)
                        """)  
    
       # ------------------- Inserindo informações ----------------------------                            
    cursor.execute (""" INSERT INTO paises(
                        Nome,
                        Nome_Oficial,
                        Capital,
                        Continente,
                        Regiao,
                        Sub_Regiao,
                        Populacao,
                        Area,
                        Fuso_Horario,
                        Bandeira,
                        Nome_Moeda,
                        Simbolo,
                        Idioma)VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",(
                        nome,
                        nomeOffi,
                        capital,
                        continente,
                        regiao,
                        subRegiao,
                        populacao,
                        area,
                        fusoH,
                        bandeira,
                        nomeMoed,
                        simbo,
                        idioma))
                                               
   
    
    print(f" \n📋 Dados adicionados no dicionário Total sobre 📌:")
    print(f"""Nome Comum: {nome} e Nome oficial: {nomeOffi}
Continente: {continente} que tem a região: {regiao} e sua subregião: {subRegiao}
Tem uma área de {area}, sua Capital é {capital} e tem como idioma {idioma}
Sua população é de {populacao} e sua moeda principal é {simbo} {nomeMoed}
Fuso horário é {fusoH} e tem sua bandeira {bandeira}\n""")
    
    dados_pais = {
        "Nome": nome,
        "Nome_Oficial": nomeOffi,
        "Capital": capital,
        "Continente": continente,
        "Regiao": regiao,
        "Sub_Regiao": subRegiao,
        "Populacao": populacao,
        "Area": area,
        "Fuso_Horario": fusoH,
        "Nome_Moeda": nomeMoed,
        "Simbolo": simbo,
        "Idioma": idioma
    }

    paises.append(dados_pais)
    conexao.commit()
    
    n -= 1
    
   
conexao.close

relatorio(paises)

print("😢 Fim das pesquisar, obrigada por participar 😁\n")
#del paises.bd - deletar banco de dados