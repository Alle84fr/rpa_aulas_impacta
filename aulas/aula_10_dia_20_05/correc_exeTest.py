import sqlite3
import requests
from docx import Document

def coletar_dados_pais(nome):
    url = f"https://restcountries.com/v3.1/name/{nome}"
    resposta = requests.get(url)

    if resposta.status_code != 200:
        print(f"Erro ao consultar o país: {nome}")
        return None

    dados = resposta.json()
    info = dados[0]

    try:
        moeda = list(info.get("currencies", {}).values())[0]
        idioma = list(info.get("languages", {}).values())[0]
    except:
        moeda = {"name": "Desconhecida", "symbol": "?"}
        idioma = "Desconhecido"

    return {
        "nome_comum": info.get("name", {}).get("common", "N/A"),
        "nome_oficial": info.get("name", {}).get("official", "N/A"),
        "capital": info.get("capital", ["N/A"])[0],
        "continente": info.get("continents", ["N/A"])[0],
        "regiao": info.get("region", "N/A"),
        "sub_regiao": info.get("subregion", "N/A"),
        "populacao": info.get("population", 0),
        "area": info.get("area", 0),
        "moeda": moeda.get("name", "N/A"),
        "simbolo_moeda": moeda.get("symbol", "?"),
        "idioma": idioma,
        "fuso_horario": info.get("timezones", ["N/A"])[0],
        "bandeira": info.get("flags", {}).get("png", "")
    }


def salvar_no_banco(dados):
    conn = sqlite3.connect("dados_paises.db")
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS paises (
            nome_comum TEXT, nome_oficial TEXT, capital TEXT,
            continente TEXT, regiao TEXT, sub_regiao TEXT,
            populacao INTEGER, area REAL, moeda TEXT,
            simbolo_moeda TEXT, idioma TEXT, fuso_horario TEXT,
            bandeira TEXT
        )
    ''')

    #cursor.execute('''
    #    INSERT INTO paises VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?,?)
    #''', tuple(dados.values()))

    conn.commit()
    conn.close()


def gerar_relatorio(lista_dados, nome_aluno="vanderson"):
    doc = Document()
    doc.add_heading("Relatório de Dados Geográficos", 0)

    for dados in lista_dados:
        doc.add_heading(dados["nome_comum"], level=1)
        for chave, valor in dados.items():
            doc.add_paragraph(f"{chave.replace('_', ' ').capitalize()}: {valor}")
        doc.add_paragraph("")

    nome_arquivo = f"relatorio_{nome_aluno}.docx"
    doc.save(nome_arquivo)
    return nome_arquivo


def main():
    lista_dados = []
    nomes = ['brazil','japan']
    for i in range(2):
        nome = nomes[i] #input(f"Digite o nome do país: ")
        dados = coletar_dados_pais(nome)
        if dados:
            salvar_no_banco(dados)
            lista_dados.append(dados)

    if lista_dados:
        nome_arquivo = gerar_relatorio(lista_dados)
        print(f"Relatorio gerado: {nome_arquivo}")


if __name__ == '__main__':
    main()